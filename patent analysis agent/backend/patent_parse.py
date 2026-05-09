"""특허 텍스트 추출, 청구항 분리, 키워드 보강, 특허용 청크 분할."""
from __future__ import annotations

import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from description_anchor import paragraph_anchor_from_text


# KR / US 공통: 청구·섹션 경계 (PRD FR-05)
# --- US: 섹션·독립/종속 청구 휴리스틱 (요약·명세·CLAIMS·인용 구문) -----------------

_RX_US_DEP_CLAIM_HEAD = re.compile(
    r"(?is)^.{0,360}?"
    r"\b(?:according\s+to|dependent\s+upon|as\s+recited\s+in|as\s+set\s+forth\s+in|"
    r"from|under)\s+claim\s*\d+|"
    r"\bthe\s+[\w\s]{0,60}?\b(?:of|according\s+to)\s+claim\s*\d+",
)

_RX_US_INDEP_BODY = re.compile(
    r"(?is)\b(?:comprising|consisting\s+essentially\s+of|consisting\s+of|including)\b",
)


def _us_claims_search_window(full_text: str) -> str:
    """
    Abstract(및 Summary of the invention) 이후만 스캔해,
    본문 앞부분에 잘못 등장한 'Claims' 단어에 속지 않도록 한다.
    """
    t = full_text.strip()
    cut = 0
    for hdr in (
        r"abstract",
        r"summary\s+of\s+the\s+invention",
        r"summary",
    ):
        m = re.search(rf"(?im)(?:^|[\n\r])\s*{hdr}\s*[^\n]*[\n\r]+", t[:20000])
        if m and m.end() > cut:
            cut = m.end()
    return t[cut:] if cut else t


def _us_snippet_invention_object_lines(full_text: str, limit: int = 2400) -> str:
    """
    Detailed description 등에서 '발명의 목적'류 문장을 잘라 BM25·디버그 보강용으로 사용할 수 있는 힌트.
    (권리범위 추출의 직접 입력은 아니고, 동일 파일 내 기술 초점 키워드 신호.)
    """
    t = full_text.strip()
    pat = (
        r"(?is)(?:^|[\n\r])\s*"
        r"(?:an\s+)?objects?\s+of\s+(?:the\s+)?(?:present\s+)?invention\s+is\s+to\s+"
        r"[^\n]{0,240}(?:[\n\r]+[^\n]{0,240}){0,8}"
    )
    parts: list[str] = []
    for m in re.finditer(pat, t[:80000]):
        parts.append(m.group(0).strip())
        if sum(len(p) for p in parts) >= limit:
            break
    blob = "\n".join(parts)
    return blob[:limit]


def _us_first_line_looks_dependent_claim(opening: str) -> bool:
    """종속항: 'of claim N' / 'according to claim N' 등 상위 항 인용이 앞부분에 있는지."""
    s = (opening or "").strip()[:420]
    if not s:
        return False
    return bool(_RX_US_DEP_CLAIM_HEAD.search(s))


def _us_block_looks_independent_claim(block: str) -> bool:
    """
    독립항 휴리스틱: 상위 항 인용이 앞부분에 없고,
    comprising / including / consisting of 등 권리범위 서술이 포함된다.
    """
    b = (block or "").strip()
    if len(b) < 30:
        return False
    head = b[:520]
    if _us_first_line_looks_dependent_claim(head):
        return False
    return bool(_RX_US_INDEP_BODY.search(b[:8000]))


def _us_enumerated_claim_blocks(claims: str) -> list[tuple[int, int, int, str]]:
    """`(번호).` 줄 시작 기준으로 청구 블록 경계 (시작, 끝, 번호, 본문)."""
    t = claims.strip()
    if not t:
        return []
    marks = list(re.finditer(r"(?m)^\s*(\d+)\s*[.．)]\s+", t))
    out: list[tuple[int, int, int, str]] = []
    for i, m in enumerate(marks):
        num = int(m.group(1))
        start = m.start()
        end = marks[i + 1].start() if i + 1 < len(marks) else len(t)
        out.append((num, start, end, t[start:end].strip()))
    return out


def _us_extract_lowest_independent_claim(claims: str) -> str | None:
    """
    번호 순으로 스캔해, 종속 서두가 아니고 독립항 패턴(comprising 등)인 최소 번호 청구를 채택.
    (PDF에서 1번이 종속으로만 잡히거나 취소된 뒤 13번이 장치·방법 독립항인 사례 대응.)
    """
    blocks = _us_enumerated_claim_blocks(claims)
    for num, _s, _e, blk in sorted(blocks, key=lambda x: x[0]):
        if _us_block_looks_independent_claim(blk):
            return blk[:50000]
    return None


def _us_claims_from_post_detailed(full_text: str) -> str:
    """Detailed Description 이후 텍스트에서 마지막 `Claims` 이후를 권리범위 후보로."""
    t = full_text.strip()
    dd = re.search(r"(?im)(?:^|[\n\r])\s*detailed\s+description\b[^\n]*[\n\r]+", t)
    if not dd:
        return ""
    after = t[dd.end() :]
    hits = list(re.finditer(r"(?is)(?:^|[\n\r])\s*claims?\b\s*", after))
    if not hits:
        return ""
    last = hits[-1]
    cand = after[last.end() :].strip()
    if len(cand) < 100:
        return ""
    if re.search(r"(?is)(?:^|[\n\r])\s*\d+\s*[.．)]\s+\S", cand[:25000]):
        return cand[:120000]
    return ""


# --- KR: 섹션·독립/종속 청구 (특허청구의 범위·제n항·인용·삭제) --------------------

_RX_KR_DEP_OPEN = re.compile(
    r"(?is)^\s*(?:"
    r"제\s*\d+\s*항\s*(?:내지|및|에서)\s*제\s*\d+\s*항|"
    r"제\s*\d+\s*항\s*에\s*있어서|"
    r"상기\s*제\s*\d+\s*항|"
    r"위\s*제\s*\d+\s*항"
    r")",
)


def _kr_block_marked_cancelled(block: str) -> bool:
    """삭제·취소·Cancel 표기가 앞부분에 있으면 해당 항은 제외."""
    head = (block or "").strip()[:400]
    return bool(
        re.search(r"(?is)\(삭제\)|\(취소\)|\bcancel(?:led)?\b|삭제\s*됨", head)
    )


def _kr_body_after_claim_label(block: str) -> str:
    """`제n항` 라벨 직후 본문만(뒤에 `.` `:` 등이 와도 제거)."""
    return re.sub(
        r"(?is)^\s*제\s*\d+\s*항\s*[.:：]?\s*",
        "",
        (block or "").strip(),
        count=1,
    ).strip()


def _kr_opening_refs_parent_claim(body_after_label: str) -> bool:
    """종속: 서두에 타 항 인용(제m항에 있어서, 상기 제m항 …)."""
    s = (body_after_label or "").strip()[:280]
    if not s:
        return False
    return bool(_RX_KR_DEP_OPEN.search(s))


def _kr_enumerate_claim_blocks(claims: str) -> list[tuple[int, int, int, str]]:
    """`제 n 항` 경계로 청구 블록 (번호, 시작, 끝, 텍스트)."""
    t = (claims or "").strip()
    if not t:
        return []
    marks = list(re.finditer(r"(?is)(?:^|[\n\r])\s*제\s*(\d+)\s*항\b", t))
    out: list[tuple[int, int, int, str]] = []
    for i, m in enumerate(marks):
        num = int(m.group(1))
        start = m.start()
        end = marks[i + 1].start() if i + 1 < len(marks) else len(t)
        out.append((num, start, end, t[start:end].strip()))
    return out


def _kr_claims_section_loose(full_text: str) -> str:
    """
    Step1: 【특허청구의 범위】·【청구항】·본문형 헤더 등으로 권리범위 구간 후보를 넓힌다.
    PDF에서 개행이 깨져 기본 패턴만으로 claims_text가 비는 경우 보완.
    """
    t = (full_text or "").strip()
    if len(t) < 120:
        return ""
    cands: list[str] = []
    end_a = r"(?=【\s*(?:요약|도면의|발명의\s*명칭|명\s*칭|배경)|(?:^|[\n\r])\s*요약\s*[：:]|\Z)"
    for pat in (
        rf"(?is)【\s*특허청구의\s*범위\s*】\s*(?P<body>[\s\S]+?){end_a}",
        rf"(?is)【\s*특허청구의범위\s*】\s*(?P<body>[\s\S]+?)(?=【|\Z)",
        rf"(?is)【\s*청구항\s*】\s*(?P<body>[\s\S]+?){end_a}",
        rf"(?is)(?:^|[\n\r])\s*특허청구의\s*범위\s*[：:\n\r]+\s*(?P<body>[\s\S]+?){end_a}",
        rf"(?is)(?:^|[\n\r])\s*청구항\s*[：:\n\r]+\s*(?P<body>[\s\S]+?){end_a}",
    ):
        m = re.search(pat, t)
        if m and (m.groupdict().get("body") or "").strip():
            chunk = m.group("body").strip()
            if len(chunk) > 80:
                cands.append(chunk[:150000])
    if not cands:
        win = t[max(0, int(len(t) * 0.42)) :]
        m2 = re.search(
            r"(?is)(제\s*1\s*항\b[\s\S]{60,12000}?)(?=제\s*2\s*항\b|【\s*요약|【\s*도면|\Z)",
            win,
        )
        if m2:
            cands.append(m2.group(1).strip()[:150000])
    return max(cands, key=len) if cands else ""


PATENT_SEPARATORS = [
    "\n\n제",
    "\n제",
    "\n\nClaim ",
    "\nClaim ",
    "\n\n【",
    "\n【",
    "\n\n## ",
    "\n\n# ",
    "\n\n",
    "\n",
    "。",
    ". ",
    " ",
    "",
]


@dataclass
class ParsedPatent:
    full_text: str
    claims_text: str
    body_text: str
    invention_title: str | None
    reference_terms: str  # BM25 보강용 짧은 문자열
    invention_claim_one: str | None  # 본 발명: 독립 제1항/Claim 1 텍스트만(비교 범위)
    application_number: str | None = None  # Phase 1: Contextual Retrieval·표기용


def _read_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp949", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def load_text_from_uploaded(name: str, data: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        from patent_pdf_load import load_pdf_text_from_bytes

        text, _label = load_pdf_text_from_bytes(data)
        return text
    if suffix == ".docx":
        try:
            import docx  # type: ignore

            import tempfile

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(data)
                path = tmp.name
            try:
                d = docx.Document(path)
                return "\n".join(p.text for p in d.paragraphs)
            finally:
                Path(path).unlink(missing_ok=True)
        except ImportError:
            raise RuntimeError("DOCX를 읽으려면 python-docx를 설치하세요.") from None
    return _read_txt(data)


def infer_jurisdiction(source_file: str, text: str) -> str:
    """
    관할 UI 없이 파일·본문으로 KR / US / unknown 추정.
    한·미 문헌을 한 번에 올릴 때 파일마다 따로 판별한다.
    """
    sample = text[:12000]
    name_lower = source_file.lower()

    kr_hits = 0
    kr_hits += len(re.findall(r"【[^】]{1,40}】", sample[:5000]))
    kr_hits += len(re.findall(r"제\s*\d+\s*항", sample))
    kr_hits += len(
        re.findall(
            r"(?:청구항|특허청구의\s*범위|발명의\s*명칭|기술\s*분야|배경\s*기술|요\s*약)",
            sample[:8000],
        )
    )
    if re.search(r"[\uac00-\ud7af]{3,}", sample):
        kr_hits += 2

    us_hits = 0
    us_hits += len(
        re.findall(
            r"(?im)^(abstract|claims?|brief\s+description\s+of\s+drawings|detailed\s+description|background)\s*$",
            sample[:8000],
        )
    )
    us_hits += len(re.findall(r"(?i)\bclaim\s*\d+\b", sample))
    us_hits += len(re.findall(r"(?i)\b(?:patent|publication)\s+no\.?", sample))
    if re.search(r"(?i)us[_-]?\d{6,}", name_lower):
        us_hits += 3

    if kr_hits >= 3 and kr_hits > us_hits * 1.1:
        return "KR"
    if us_hits >= 2 and us_hits > kr_hits * 1.1:
        return "US"
    return "unknown"


def extract_application_number(text: str, jurisdiction: str) -> str | None:
    """명세서 앞부분에서 KR·US 출원·공보 번호 후보를 추출."""
    t = (text or "").strip()[:80000]
    if not t:
        return None
    j = (jurisdiction or "unknown").upper()

    def _kr() -> str | None:
        for pat in (
            r"(?:출원번호|공개번호|등록번호|공보번호)\s*[：:]\s*(\d{2}-\d{4}-\d{7,}|\d{10,15})",
            r"(?<![\d-])(\d{2}-\d{4}-\d{7,})(?![\d-])",
        ):
            m = re.search(pat, t[:20000])
            if m:
                return m.group(1).strip()[:80]
        return None

    def _us() -> str | None:
        for pat in (
            r"(?i)(?:publication|patent)\s+no\.?\s*[：:#]?\s*(US\s*[\d,/]+\s*[A-Z]?\d*)",
            r"(?i)\b(US\s*\d{4}/\d{6,7}\s*[A-Z]\d?)\b",
            r"(?i)\b(US\s*\d{6,}\s*[A-Z]\d?)\b",
        ):
            m = re.search(pat, t[:25000])
            if m:
                return re.sub(r"\s+", " ", m.group(1).strip())[:80]
        return None

    if j == "KR":
        return _kr()
    if j == "US":
        return _us()
    if re.search(r"【\s*발명의\s*명칭|특허청구의\s*범위|제\s*\d+\s*항", t[:6000]):
        return _kr() or _us()
    if re.search(r"(?i)\b(?:claims?|abstract|us\s*\d)\b", t[:6000]):
        return _us() or _kr()
    return _kr() or _us()


def extract_invention_title(text: str, jurisdiction: str) -> str | None:
    patterns = [
        r"【\s*발명의\s*명칭\s*】\s*([^\n]+)",
        r"(?i)title\s*of\s*invention\s*[:\s]+([^\n]+)",
        r"(?i)^title\s*[:\s]+([^\n]+)",
    ]
    for p in patterns:
        m = re.search(p, text[:8000], re.MULTILINE)
        if m:
            return m.group(1).strip()[:500]
    return None


def extract_reference_terms_block(text: str) -> str:
    """부호의 명칭·도면 설명 등 — BM25용 토큰 나열."""
    blocks: list[str] = []
    for pat in (
        r"【\s*도면의\s*간단한\s*설명\s*】([\s\S]{0,8000}?)(?=【|$)",
        r"【\s*부호의\s*설명\s*】([\s\S]{0,8000}?)(?=【|$)",
        r"(?i)brief\s+description\s+of\s+drawings([\s\S]{0,8000}?)(?=\n[A-Z]{3,}|$)",
    ):
        m = re.search(pat, text)
        if m:
            blocks.append(m.group(1).strip()[:6000])
    # 짧은 "숫자: 용어" 라인들
    lines = re.findall(r"^\s*(\d{1,4})\s*[:：\-—]\s*([^\n]{1,120})", text, re.MULTILINE)
    if lines:
        blocks.append(" ".join(f"{a}:{b}" for a, b in lines[:80]))
    obj = _us_snippet_invention_object_lines(text)
    if obj.strip():
        blocks.append(f"INV_OBJECT_HINT: {obj.strip()[:1800]}")
    return " ".join(blocks)[:8000]


def _extract_claim_one_kr(claims: str) -> str | None:
    """
    KR 독립 제1항(또는 삭제·종속으로 1항을 쓸 수 없을 때 최소 번호 독립항).

    Step2: `제n항` 직후에 타 항 인용(제m항에 있어서 등)이 오면 종속으로 본다.
    Step3: (삭제)/Cancel 등은 제외.
    """
    t = (claims or "").strip()
    if not t:
        return None
    blocks = _kr_enumerate_claim_blocks(t)
    if not blocks:
        m = re.search(
            r"(?is)(제\s*1\s*항\b)([\s\S]+?)(?=제\s*2\s*항\b|\Z)",
            t,
        )
        if m:
            blocks = [(1, m.start(1), m.end(0), m.group(0).strip())]
        else:
            return None

    by_num: dict[int, str] = {}
    for num, _s, _e, blk in blocks:
        by_num[num] = blk

    def _pick_independent_for_num(n: int) -> str | None:
        blk = by_num.get(n)
        if not blk or _kr_block_marked_cancelled(blk):
            return None
        inner = _kr_body_after_claim_label(blk)
        if _kr_opening_refs_parent_claim(inner):
            return None
        if len(blk.strip()) < 20:
            return None
        return blk.strip()[:50000]

    if 1 in by_num:
        got = _pick_independent_for_num(1)
        if got:
            return got
    for n in sorted(by_num):
        got = _pick_independent_for_num(n)
        if got:
            return got
    return None


def _extract_claim_one_us(claims: str) -> str | None:
    """
    US: 번호 청구 `1.` … (또는 `Claim 1`) 블록.

    종속항은 앞부분에 `of claim N` / `according to claim N` 등이 오는 경우가 많아,
    1번이 종속으로만 인식되면 **comprising·including** 등 독립항 패턴이 있는 더 낮은 번호
    독립항(예: 장치 1항·방법 13항)을 순차 스캔한다.
    """
    t = claims.strip()
    if not t:
        return None
    candidates: list[str] = []

    def _push(block: str | None) -> None:
        if not block:
            return
        b = block.strip()
        if len(b) > 12:
            candidates.append(b[:50000])

    if re.match(r"(?is)^\s*(?:claim\s*)?1\s*[.．)]\s+\S", t):
        m0 = re.search(
            r"(?is)^\s*((?:claim\s*)?1\s*[.．)]\s+[\s\S]*?)(?=(?:^|[\n\r])\s*(?:claim\s*)?2\s*[.．)]\s+|[\n\r]\s*claim\s+2\b|\Z)",
            t,
        )
        if m0:
            _push(m0.group(1))
    m = re.search(
        r"(?is)((?:^|[\n\r])\s*(?:claim\s*)?1\s*[.．)]\s+[\s\S]*?)(?=[\n\r]\s*(?:claim\s*)?2\s*[.．)]|[\n\r]\s*claim\s+2\b|\Z)",
        t,
    )
    if m:
        _push(m.group(1))
    m2 = re.search(
        r"(?is)((?:^|[\n\r])\s*claim\s+1\s*[.:]?\s*[\s\S]*?)(?=[\n\r]\s*claim\s+2\b|\Z)",
        t,
    )
    if m2:
        _push(m2.group(1))

    for cand in candidates:
        if not _us_first_line_looks_dependent_claim(cand[:520]):
            return cand
    if candidates:
        alt = _us_extract_lowest_independent_claim(t)
        if alt:
            return alt
        return candidates[0]
    return _us_extract_lowest_independent_claim(t)


def _us_claims_block_loose(full_text: str) -> str:
    """
    US PDF 텍스트에서 `Claims` 헤더 직후 줄바꿈이 없거나(예: CLAIMS 1.),
    본문 중간에 잘못 잡힌 Claims 때문에 기본 정규식이 빈 문자열을 줄 때 보완.

    - 문서 **끝부근**의 마지막 `Claims`(또는 what is claimed) 이후 ~끝까지를 후보로 삼는다.
    - 미국 공보·등록 특허는 청구항이 문서 후반에 오는 경우가 많다.
    """
    t = (full_text or "").strip()
    if len(t) < 200:
        return ""
    # "What is claimed is:" / "We claim:" 등 (헤더 없이 바로 번호 청구)
    for pat in (
        r"(?is)\bwhat\s+is\s+claimed\s+is\s*[\n\r:]*\s*",
        r"(?is)\bwe\s+claim\s*:\s*[\n\r]*",
    ):
        m = re.search(pat, t)
        if m:
            cand = t[m.end() :].strip()
            if len(cand) > 120 and re.search(r"(?is)(?:^|[\n\r])\s*(?:claim\s*)?1\s*[.．)]\s+", cand[:8000]):
                return cand[:120000]
    win = _us_claims_search_window(t)
    hits = list(re.finditer(r"(?is)(?:^|[\n\r])\s*claims?\b\s*", win))
    if not hits:
        return ""
    last = hits[-1]
    cand = win[last.end() :].strip()
    if len(cand) < 80:
        return ""
    # "Claims" 직후가 바로 `1.` 인 경우(한 줄 붙음)
    if re.match(r"(?is)^(?:claim\s*)?1\s*[.．)]\s+\S", cand):
        return cand[:120000]
    # 일반: Claims 다음 줄부터
    if re.search(r"(?is)(?:^|[\n\r])\s*(?:claim\s*)?1\s*[.．)]\s+", cand[:20000]):
        return cand[:120000]
    return ""


def _us_claim1_tail_fallback(full_text: str) -> str:
    """
    `Claims` 표기가 PDF 추출에서 사라졌거나 본문 `1.` 만 남은 경우,
    문서 **후반**(요약·명세 이후)에서 `1. ... 2.` 패턴으로 청구 블록을 추정한다.
    """
    t = (full_text or "").strip()
    surf = _us_claims_search_window(t)
    if len(surf) < 600:
        return ""
    start = int(len(surf) * 0.12)
    tail = surf[start:]
    m = re.search(
        r"(?is)((?:^|[\n\r])\s*(?:claim\s*)?1\s*[.．)]\s+[\s\S]+?)(?=(?:^|[\n\r])\s*(?:claim\s*)?2\s*[.．)]\s+)",
        tail,
    )
    if not m:
        return ""
    block = m.group(1).strip()
    if len(block) < 40:
        return ""
    return block[:120000]


def _regex_split_limitations_us(text: str) -> list[str]:
    """
    US Claim 1: preamble·comprising 블록과 각 wherein 절을 분리.
    wherein 앞은 쉼표/세미콜론/개행 중 하나(특허 문장 관행).
    """
    t = text.strip()
    if not t:
        return []
    # 마침표·괄호 뒤 줄바꿈 wherein (US 출원 문장 관행)
    split_rx = r"(?is)(?<=[,;，；\n\r.．\)\]］）〉])\s*\bwherein\b"
    matches = list(re.finditer(split_rx, t))
    if not matches:
        return [t]
    out: list[str] = []
    head = t[: matches[0].start()].strip()
    if head:
        out.append(head)
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(t)
        seg = t[start:end].strip()
        if seg:
            out.append(seg)
    return out


def _regex_split_limitations_kr(text: str) -> list[str]:
    """
    KR 제1항: 영문 wherein이 섞인 출원은 US 규칙 우선.
    그 외는 '에 있어서'(의존항 도입구) 또는 마침표 뒤 '상기' 도입 등으로 1차 분할 시도.
    """
    t = text.strip()
    if not t:
        return []
    if re.search(r"(?is)\bwherein\b", t):
        return _regex_split_limitations_us(t)
    # 마침표(다.) 뒤 이어지는 한정(상기…)
    pieces = re.split(r"(?<=다\.)\s+(?=상기)", t)
    if len(pieces) > 1:
        return [p.strip() for p in pieces if p.strip()]
    return [t]


def regex_split_claim_limitations(claim_text: str, jurisdiction: str) -> list[str]:
    """
    Claim 1 / 제1항을 한정 단위로 1차 분할(정규식).
    US wherein 각각은 별도 세그먼트가 되도록 한다.
    """
    t = (claim_text or "").strip()
    if not t:
        return []
    j = (jurisdiction or "unknown").upper()
    if j == "KR":
        segs = _regex_split_limitations_kr(t)
        if len(segs) > 1:
            return segs
        segs_us = _regex_split_limitations_us(t)
        return segs_us if len(segs_us) > 1 else [t]
    if j == "US":
        segs = _regex_split_limitations_us(t)
        return segs if len(segs) > 1 else [t]
    segs_us = _regex_split_limitations_us(t)
    if len(segs_us) > 1:
        return segs_us
    segs_kr = _regex_split_limitations_kr(t)
    return segs_kr if len(segs_kr) > 1 else [t]


def extract_invention_claim_one(claims_text: str, jurisdiction: str) -> str | None:
    """
    본 발명 비교 범위: 독립 제1항(KR) / Claim 1(US) 텍스트만.
    관할 unknown 시 KR 패턴 후 US 패턴 순으로 시도.
    """
    if not claims_text or len(claims_text.strip()) < 5:
        return None
    j = (jurisdiction or "unknown").upper()
    if j == "KR":
        return _extract_claim_one_kr(claims_text) or _extract_claim_one_us(claims_text)
    if j == "US":
        return _extract_claim_one_us(claims_text) or _extract_claim_one_kr(claims_text)
    return _extract_claim_one_kr(claims_text) or _extract_claim_one_us(claims_text)


def claim1_or_claims_section_fallback(claims_text: str | None, jurisdiction: str) -> str:
    """
    세션 `claim1_full`·비교 표·Phase2 입력용 문자열.

    우선 제1항/Claim 1만 추출하고, 실패해도 **청구항 구간**이 충분히 있으면 그 앞부분을
    사용한다. PDF 줄바꿈·표기 때문에 정규식이 제1항을 못 잡는 경우가 많아, 이 경로가
    비어 있으면 비교 표가 \"청구 텍스트 없음\"으로 끊기지만 벡터 청크 수는 여전히 클 수 있다.
    """
    claims = (claims_text or "").strip()
    if not claims:
        return ""
    first = extract_invention_claim_one(claims, jurisdiction)
    if (first or "").strip():
        return (first or "").strip()
    if len(claims) > 80:
        return claims[:20000].strip()
    return ""


def extract_claim_preamble_plain(claim_text: str) -> str:
    """
    청구항 설명 표 첫 행용 preamble 텍스트(플레인).

    US: 번호 라벨 제거 후 ``… comprising[:;]`` 까지(포함).
    KR: ``… 특징으로 하는`` 까지(포함) 우선, 없으면 첫 줄 일부.
    """
    raw = (claim_text or "").strip()
    if not raw:
        return ""
    head = raw[:4000]
    t = re.sub(
        r"^\s*(?:claim\s*)?(\d{1,3})\s*[.．:]\s*",
        "",
        head,
        count=1,
        flags=re.I,
    ).strip()
    m_us = re.search(r"(?is)^(.{10,1200}?\bcomprising\s*[;:]\s*)", t[:2500])
    if m_us:
        return m_us.group(1).strip()
    m_kr = re.search(r"(?is)^(.{10,1200}?(?:을\s*)?특징으로\s*하는)", t[:2500])
    if m_kr:
        return m_kr.group(1).strip()
    return t.split("\n", 1)[0].strip()[:400]


def claim_mapping_claim_column_label(jurisdiction: str) -> str:
    """청구-명세 매핑 표 1열 헤더: KR 공보는 청구항, 그 외(US 등)는 Claim."""
    j = (jurisdiction or "").strip().upper()
    if j == "KR":
        return "청구항"
    return "Claim"


def _kr_collect_independent_claim_blocks(claims: str) -> list[tuple[int, str]]:
    """제n항 블록 중 종속·삭제 제외한 독립항 (번호, 전문)."""
    t = (claims or "").strip()
    if not t:
        return []
    blocks = _kr_enumerate_claim_blocks(t)
    if not blocks:
        m = re.search(
            r"(?is)(제\s*1\s*항\b)([\s\S]+?)(?=제\s*2\s*항\b|\Z)",
            t,
        )
        if m:
            blocks = [(1, m.start(1), m.end(0), m.group(0).strip())]
        else:
            return []
    by_num: dict[int, str] = {}
    for num, _s, _e, blk in blocks:
        by_num[num] = blk
    out: list[tuple[int, str]] = []
    for n in sorted(by_num):
        blk = by_num.get(n)
        if not blk or _kr_block_marked_cancelled(blk):
            continue
        inner = _kr_body_after_claim_label(blk)
        if _kr_opening_refs_parent_claim(inner):
            continue
        if len(blk.strip()) < 20:
            continue
        out.append((n, blk.strip()[:50000]))
    return out


def _us_collect_independent_claim_blocks(claims: str) -> list[tuple[int, str]]:
    """번호 청구 블록 중 독립항 휴리스틱 통과 항목."""
    blocks = _us_enumerated_claim_blocks(claims)
    out: list[tuple[int, str]] = []
    for num, _s, _e, blk in sorted(blocks, key=lambda x: x[0]):
        if _us_block_looks_independent_claim(blk):
            out.append((num, blk.strip()[:50000]))
    return out


def extract_all_independent_claims(claims_text: str, jurisdiction: str) -> list[dict[str, Any]]:
    """
    본 발명 청구항 구간에서 독립항 전부.

    각 항목: ``claim_num``, ``text``(최대 50k), ``preview``(UI용 앞부분).
    """
    j = (jurisdiction or "unknown").upper()
    pairs: list[tuple[int, str]] = []
    if j == "KR":
        pairs = _kr_collect_independent_claim_blocks(claims_text)
        if not pairs:
            pairs = _us_collect_independent_claim_blocks(claims_text)
    elif j == "US":
        pairs = _us_collect_independent_claim_blocks(claims_text)
        if not pairs:
            pairs = _kr_collect_independent_claim_blocks(claims_text)
    else:
        pairs = _kr_collect_independent_claim_blocks(claims_text)
        if not pairs:
            pairs = _us_collect_independent_claim_blocks(claims_text)
    out: list[dict[str, Any]] = []
    for n, txt in pairs:
        pv = txt if len(txt) <= 220 else txt[:217].rstrip() + "…"
        out.append({"claim_num": n, "text": txt, "preview": pv})
    return out


def invention_claim_one_warnings(source_file: str, parsed: ParsedPatent) -> list[str]:
    """제1항 추출 실패 시 인덱스 빌드 단계에서 안내."""
    if parsed.invention_claim_one:
        return []
    if len((parsed.claims_text or "").strip()) < 30:
        return []
    return [
        f"[본 발명] `{source_file}`: 청구항 구간은 있으나 **제1항 / Claim 1**을 자동 분리하지 못했습니다. "
        "비교 표·메타데이터에는 청구항 구간 앞부분을 임시로 쓸 수 있습니다. PDF 텍스트·줄바꿈을 정리하면 독립항만 추출되는 경우가 많습니다."
    ]


def split_claims_and_body(text: str, jurisdiction: str) -> tuple[str, str]:
    """청구항 구간 추출; 실패 시 전체를 본문으로."""
    t = text.strip()
    claims = ""

    if jurisdiction.upper() == "US" or re.search(r"(?i)\bclaims?\b", t[:5000]):
        m = re.search(
            r"(?is)(^(?:claims?)\s*$|\n\s*(?:claims?)\s*\n)([\s\S]+?)(?=\n\s*(?:abstract|description|detailed|background)\b|\Z)",
            t,
        )
        if m:
            claims = m.group(2).strip()
        if not claims:
            m_glued = re.search(
                r"(?is)(?:^|[\n\r])\s*claims?\s+(?!is\b)([\s\S]+?)(?=\n\s*(?:abstract|description|detailed|background)\b|\Z)",
                t,
            )
            if m_glued:
                claims = m_glued.group(1).strip()
        if not claims and (jurisdiction.upper() == "US" or re.search(r"(?is)\bclaims?\b", t)):
            claims = _us_claims_block_loose(t)
        if not claims and jurisdiction.upper() == "US":
            claims = _us_claims_from_post_detailed(t)
        if not claims and jurisdiction.upper() == "US":
            claims = _us_claim1_tail_fallback(t)

    if not claims:
        m = re.search(
            r"(【\s*(?:청구항|특허청구의\s*범위)\s*】|【\s*특허청구의범위\s*】|"
            r"^\s*청구항\s*$|\n\s*청구항\s*\n)([\s\S]+?)(?=【\s*(?:요약|abstract|도면|발명의\s*명칭)|\Z)",
            t,
            re.MULTILINE,
        )
        if m:
            claims = m.group(2).strip()

    if (not claims or len(claims) < 100) and jurisdiction.upper() in ("KR", "UNKNOWN"):
        kr_loose = _kr_claims_section_loose(t)
        if len(kr_loose) > len(claims or ""):
            claims = kr_loose

    if claims and len(claims) > 50:
        body = t.replace(claims, "", 1).strip()
        if len(body) < 200:
            body = t
        return claims, body
    return "", t


def parse_patent_document(raw: str, jurisdiction: str, source_file: str) -> ParsedPatent:
    claims, body = split_claims_and_body(raw, jurisdiction)
    title = extract_invention_title(raw, jurisdiction)
    ref_terms = extract_reference_terms_block(raw)
    claim_one = extract_invention_claim_one(claims, jurisdiction) if claims else None
    app_no = extract_application_number(raw, jurisdiction)
    return ParsedPatent(
        full_text=raw,
        claims_text=claims,
        body_text=body or raw,
        invention_title=title,
        reference_terms=ref_terms,
        invention_claim_one=claim_one,
        application_number=app_no,
    )


def patent_text_splitter() -> RecursiveCharacterTextSplitter:
    from config import CHUNK_OVERLAP, CHUNK_SIZE

    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=PATENT_SEPARATORS,
        length_function=len,
    )


def augment_chunk_for_bm25(page: str, title: str | None, ref_terms: str) -> str:
    parts = []
    if title:
        parts.append(f"TITLE: {title}")
    if ref_terms:
        parts.append(f"REF_TERMS: {ref_terms[:4000]}")
    parts.append(page)
    return "\n".join(parts)


def documents_from_parsed(
    parsed: ParsedPatent,
    *,
    doc_type_base: str,
    jurisdiction: str,
    source_file: str,
    session_id: str,
) -> list[Document]:
    """invention / prior_art 및 *_claims Document 리스트."""
    splitter = patent_text_splitter()
    docs: list[Document] = []
    uploaded_at = datetime.now(timezone.utc).isoformat()
    app_meta = (parsed.application_number or "").strip()
    base_meta = {
        "jurisdiction": jurisdiction,
        "source_file": source_file,
        "session_id": session_id,
        "uploaded_at": uploaded_at,
        "metadata_schema_version": "v1",
        "application_number": app_meta,
    }

    def add_chunks(
        text: str,
        dtype: str,
        section_type: str,
        *,
        extra_meta: dict | None = None,
    ) -> None:
        if not text.strip():
            return
        chunks = splitter.split_text(text)
        for i, chunk in enumerate(chunks):
            cid = f"{session_id}_{source_file}_{dtype}_{section_type}_{i}_{uuid.uuid4().hex[:8]}"
            meta = {
                **base_meta,
                "doc_type": dtype,
                "section_type": section_type,
                "chunk_id": cid,
                "invention_title": parsed.invention_title or "",
                **(extra_meta or {}),
            }
            if section_type == "body" and dtype in ("invention", "prior_art"):
                pa = paragraph_anchor_from_text(chunk)
                if pa:
                    meta["paragraph_anchor"] = pa
            bm25_page = augment_chunk_for_bm25(
                chunk, parsed.invention_title, parsed.reference_terms
            )
            docs.append(Document(page_content=bm25_page, metadata=meta))

    if doc_type_base == "invention":
        # 비교 범위: 독립 제1항/Claim 1만 인덱싱(제2항 이하·다른 독립항은 NFR).
        if parsed.invention_claim_one:
            add_chunks(
                parsed.invention_claim_one,
                "invention_claims",
                "claims",
                extra_meta={
                    "claim_scope": "independent_claim_1",
                    "claim_number": 1,
                    "independent_claim": True,
                },
            )
        if parsed.body_text:
            add_chunks(
                parsed.body_text,
                "invention",
                "body",
                extra_meta={
                    "claim_scope": "description_body",
                    "claim_number": 0,
                    "independent_claim": False,
                },
            )
    else:
        if parsed.claims_text:
            add_chunks(
                parsed.claims_text,
                "prior_art_claims",
                "claims",
                extra_meta={"claim_number": 0, "independent_claim": False},
            )
        if parsed.body_text:
            add_chunks(
                parsed.body_text,
                "prior_art",
                "body",
                extra_meta={"claim_number": 0, "independent_claim": False},
            )

    return docs
