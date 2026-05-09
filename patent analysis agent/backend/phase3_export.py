"""Markdown 텍스트를 단순 DOCX 단락으로 내보내기 (표는 평문 줄로 저장)."""
from __future__ import annotations

import io
import re


def markdown_report_to_docx_bytes(markdown_text: str) -> bytes:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise RuntimeError("DOCX 내보내기에는 python-docx가 필요합니다.") from e

    doc = docx.Document()
    doc.add_heading("Patent Analysis Agent — Report", level=1)
    for block in re.split(r"\n{2,}", markdown_text.strip() or "(empty)"):
        for line in block.split("\n"):
            doc.add_paragraph(line.rstrip() or " ")
        doc.add_paragraph("")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
